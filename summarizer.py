import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from deep_translator import GoogleTranslator
from gtts import gTTS
#import playsound
import tempfile
import os
import threading
from docx import Document
import PyPDF2
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.summarizers.lsa import LsaSummarizer
from sumy.summarizers.luhn import LuhnSummarizer
from sumy.summarizers.kl import KLSummarizer
import time
import nltk
from nltk.tokenize import sent_tokenize
import json
from datetime import datetime
import hashlib

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

# Language dictionary
languages = {
    'English': 'en', 'Hindi': 'hi', 'Bengali': 'bn', 'Telugu': 'te',
    'Marathi': 'mr', 'Tamil': 'ta', 'Gujarati': 'gu', 'Malayalam': 'ml',
    'Kannada': 'kn', 'Punjabi': 'pa', 'Urdu': 'ur', 'Spanish': 'es',
    'French': 'fr', 'German': 'de', 'Italian': 'it', 'Russian': 'ru',
    'Chinese (Simplified)': 'zh-CN', 'Japanese': 'ja', 'Korean': 'ko',
    'Arabic': 'ar', 'Portuguese': 'pt', 'Dutch': 'nl'
}


class ProfessionalSummarizer:
    def __init__(self, root):
        self.root = root
        self.root.title("🎯 Professional Text Summarization Suite v3.0")
        self.root.geometry("1300x850")
        self.root.configure(bg='#f5f5f5')

        # Variables
        self.uploaded_content = ""
        self.history = []
        self.favorites = []
        self.current_theme = "light"

        # Create GUI
        self.create_menu_bar()
        self.create_main_interface()
        self.create_status_bar()

        # Load saved data
        self.load_settings()
        self.load_history()
        self.load_favorites()

        # Play welcome
        self.play_welcome()

    def create_menu_bar(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)

        # File Menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="📁 File", menu=file_menu)
        file_menu.add_command(label="📂 New Session", command=self.new_session)
        file_menu.add_command(label="📄 Open File", command=self.upload_document)
        file_menu.add_command(label="💾 Save Summary", command=self.save_summary)
        file_menu.add_separator()
        file_menu.add_command(label="🚪 Exit", command=self.root.quit)

        # Tools Menu
        tools_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="🛠️ Tools", menu=tools_menu)
        tools_menu.add_command(label="🌐 Translator", command=self.open_translator)
        tools_menu.add_command(label="📊 Text Analyzer", command=self.open_analyzer)
        tools_menu.add_command(label="🔄 Batch Processing", command=self.open_batch_processor)
        tools_menu.add_separator()
        tools_menu.add_command(label="⭐ Favorites", command=self.open_favorites)
        tools_menu.add_command(label="📜 History", command=self.open_history)

        # Settings Menu
        settings_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="⚙️ Settings", menu=settings_menu)
        settings_menu.add_command(label="🎨 Change Theme", command=self.change_theme)
        settings_menu.add_command(label="🔊 Audio Settings", command=self.audio_settings)
        settings_menu.add_separator()
        settings_menu.add_command(label="🗑️ Clear Cache", command=self.clear_cache)

        # Help Menu
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="❓ Help", menu=help_menu)
        help_menu.add_command(label="📖 Documentation", command=self.show_documentation)
        help_menu.add_command(label="ℹ️ About", command=self.show_about)

    def create_main_interface(self):
        # Header
        header_frame = tk.Frame(self.root, bg='#2c3e50', height=80)
        header_frame.pack(fill=tk.X)
        header_frame.pack_propagate(False)

        title_label = tk.Label(header_frame, text="🎯 PROFESSIONAL TEXT SUMMARIZATION SUITE",
                               font=("Arial", 22, "bold"), bg='#2c3e50', fg='white')
        title_label.pack(pady=20)

        # Main container
        main_container = tk.Frame(self.root, bg='#f5f5f5')
        main_container.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        # Left Panel - Controls
        left_panel = tk.Frame(main_container, bg='#f5f5f5', width=350)
        left_panel.pack(side=tk.LEFT, fill=tk.Y, padx=(0, 10))
        left_panel.pack_propagate(False)

        # File Operations Frame
        file_frame = tk.LabelFrame(left_panel, text="📁 File Operations",
                                   font=("Arial", 12, "bold"), bg='#f5f5f5')
        file_frame.pack(fill=tk.X, pady=(0, 10))

        tk.Button(file_frame, text="📤 Upload Document", command=self.upload_document,
                  bg='#3498db', fg='white', font=("Arial", 11), height=2).pack(fill=tk.X, padx=10, pady=5)
        tk.Button(file_frame, text="📋 Paste from Clipboard", command=self.paste_text,
                  bg='#9b59b6', fg='white', font=("Arial", 11), height=2).pack(fill=tk.X, padx=10, pady=5)
        tk.Button(file_frame, text="🗑️ Clear All", command=self.clear_all,
                  bg='#e74c3c', fg='white', font=("Arial", 11), height=2).pack(fill=tk.X, padx=10, pady=5)

        # Settings Frame
        settings_frame = tk.LabelFrame(left_panel, text="⚙️ Summary Settings",
                                       font=("Arial", 12, "bold"), bg='#f5f5f5')
        settings_frame.pack(fill=tk.X, pady=(0, 10))

        # Number of sentences
        sent_frame = tk.Frame(settings_frame, bg='#f5f5f5')
        sent_frame.pack(fill=tk.X, padx=10, pady=5)
        tk.Label(sent_frame, text="📊 Number of Sentences:", font=("Arial", 10),
                 bg='#f5f5f5').pack(side=tk.LEFT)
        self.sentence_count = tk.Entry(sent_frame, width=5, font=("Arial", 10))
        self.sentence_count.pack(side=tk.RIGHT)
        self.sentence_count.insert(0, "5")

        # Summary style
        style_frame = tk.Frame(settings_frame, bg='#f5f5f5')
        style_frame.pack(fill=tk.X, padx=10, pady=5)
        tk.Label(style_frame, text="🎨 Summary Style:", font=("Arial", 10),
                 bg='#f5f5f5').pack(side=tk.LEFT)
        self.summary_style = ttk.Combobox(style_frame, values=['📝 Concise', '📖 Detailed', '⚖️ Balanced'],
                                          width=15)
        self.summary_style.pack(side=tk.RIGHT)
        self.summary_style.set('⚖️ Balanced')

        # Language for output
        lang_frame = tk.Frame(settings_frame, bg='#f5f5f5')
        lang_frame.pack(fill=tk.X, padx=10, pady=5)
        tk.Label(lang_frame, text="🌐 Output Language:", font=("Arial", 10),
                 bg='#f5f5f5').pack(side=tk.LEFT)
        self.output_lang = ttk.Combobox(lang_frame, values=list(languages.keys()), width=15)
        self.output_lang.pack(side=tk.RIGHT)
        self.output_lang.set('English')

        # Summarization Methods Frame - MAIN BUTTONS
        methods_frame = tk.LabelFrame(left_panel, text="🎯 SUMMARIZATION METHODS",
                                      font=("Arial", 13, "bold"), bg='#f5f5f5', fg='#2c3e50')
        methods_frame.pack(fill=tk.X, pady=(0, 10))

        # Button 1: Most Important Sentences (LexRank)
        btn1_frame = tk.Frame(methods_frame, bg='#f5f5f5')
        btn1_frame.pack(fill=tk.X, padx=10, pady=5)
        btn1 = tk.Button(btn1_frame, text="🎯 MOST IMPORTANT SENTENCES",
                         command=self.lexrank_summary,
                         bg='#27ae60', fg='white', font=("Arial", 11, "bold"),
                         height=2, cursor="hand2")
        btn1.pack(fill=tk.X)
        tk.Label(btn1_frame, text="📌 LexRank - Finds central & most representative sentences",
                 font=("Arial", 8), bg='#f5f5f5', fg='#666').pack()

        # Button 2: Main Topic Summary (LSA)
        btn2_frame = tk.Frame(methods_frame, bg='#f5f5f5')
        btn2_frame.pack(fill=tk.X, padx=10, pady=5)
        btn2 = tk.Button(btn2_frame, text="🔬 MAIN TOPIC SUMMARY",
                         command=self.lsa_summary,
                         bg='#2980b9', fg='white', font=("Arial", 11, "bold"),
                         height=2, cursor="hand2")
        btn2.pack(fill=tk.X)
        tk.Label(btn2_frame, text="📌 LSA - Identifies and focuses on main themes/topics",
                 font=("Arial", 8), bg='#f5f5f5', fg='#666').pack()

        # Button 3: Quick Key Points (Luhn)
        btn3_frame = tk.Frame(methods_frame, bg='#f5f5f5')
        btn3_frame.pack(fill=tk.X, padx=10, pady=5)
        btn3 = tk.Button(btn3_frame, text="⚡ QUICK KEY POINTS",
                         command=self.luhn_summary,
                         bg='#e67e22', fg='white', font=("Arial", 11, "bold"),
                         height=2, cursor="hand2")
        btn3.pack(fill=tk.X)
        tk.Label(btn3_frame, text="📌 Luhn - Extracts important keywords and key phrases",
                 font=("Arial", 8), bg='#f5f5f5', fg='#666').pack()

        # Button 4: Detailed Important Summary (KL)
        btn4_frame = tk.Frame(methods_frame, bg='#f5f5f5')
        btn4_frame.pack(fill=tk.X, padx=10, pady=5)
        btn4 = tk.Button(btn4_frame, text="📚 DETAILED IMPORTANT SUMMARY",
                         command=self.kl_summary,
                         bg='#8e44ad', fg='white', font=("Arial", 11, "bold"),
                         height=2, cursor="hand2")
        btn4.pack(fill=tk.X)
        tk.Label(btn4_frame, text="📌 KL - Preserves maximum information with full details",
                 font=("Arial", 8), bg='#f5f5f5', fg='#666').pack()

        # Action Buttons Frame
        action_frame = tk.LabelFrame(left_panel, text="🎬 Actions",
                                     font=("Arial", 12, "bold"), bg='#f5f5f5')
        action_frame.pack(fill=tk.X)

        btn_frame1 = tk.Frame(action_frame, bg='#f5f5f5')
        btn_frame1.pack(fill=tk.X, padx=10, pady=5)
        tk.Button(btn_frame1, text="🔄 Compare All Methods", command=self.compare_all,
                  bg='#f39c12', fg='white', font=("Arial", 10, "bold")).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame1, text="🔊 Listen to Summary", command=self.text_to_speech,
                  bg='#1abc9c', fg='white', font=("Arial", 10, "bold")).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame1, text="⭐ Add to Favorites", command=self.add_to_favorites,
                  bg='#e74c3c', fg='white', font=("Arial", 10, "bold")).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame1, text="📥 Export Summary", command=self.export_summary,
                  bg='#34495e', fg='white', font=("Arial", 10, "bold")).pack(fill=tk.X, pady=2)

        # Right Panel - Text Areas
        right_panel = tk.Frame(main_container, bg='#f5f5f5')
        right_panel.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)

        # Input Text Area
        input_frame = tk.LabelFrame(right_panel, text="📄 INPUT TEXT",
                                    font=("Arial", 12, "bold"), bg='#f5f5f5', fg='#2c3e50')
        input_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        self.input_text = scrolledtext.ScrolledText(input_frame, wrap=tk.WORD,
                                                    font=("Arial", 11), height=12,
                                                    bg='white', fg='black',
                                                    selectbackground='#3498db')
        self.input_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Word and character count for input
        input_info = tk.Frame(input_frame, bg='#f5f5f5')
        input_info.pack(fill=tk.X, padx=10, pady=(0, 5))
        self.input_word_count = tk.Label(input_info, text="Words: 0 | Characters: 0",
                                         font=("Arial", 9), bg='#f5f5f5', fg='#666')
        self.input_word_count.pack(side=tk.LEFT)
        self.input_text.bind('<KeyRelease>', self.update_word_count)

        # Output Text Area
        output_frame = tk.LabelFrame(right_panel, text="✨ SUMMARY OUTPUT",
                                     font=("Arial", 12, "bold"), bg='#f5f5f5', fg='#2c3e50')
        output_frame.pack(fill=tk.BOTH, expand=True)

        self.output_text = scrolledtext.ScrolledText(output_frame, wrap=tk.WORD,
                                                     font=("Arial", 11), height=12,
                                                     bg='#fffef7', fg='black')
        self.output_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Output info
        output_info = tk.Frame(output_frame, bg='#f5f5f5')
        output_info.pack(fill=tk.X, padx=10, pady=(0, 5))
        self.output_method = tk.Label(output_info, text="Method: None",
                                      font=("Arial", 9, "bold"), bg='#f5f5f5', fg='#27ae60')
        self.output_method.pack(side=tk.LEFT)

    def create_status_bar(self):
        self.status_bar = tk.Frame(self.root, bg='#34495e', height=30)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)

        self.status_label = tk.Label(self.status_bar, text="✅ Ready",
                                     font=("Arial", 9), bg='#34495e', fg='white')
        self.status_label.pack(side=tk.LEFT, padx=10, pady=5)

        self.time_label = tk.Label(self.status_bar, text="",
                                   font=("Arial", 9), bg='#34495e', fg='white')
        self.time_label.pack(side=tk.RIGHT, padx=10, pady=5)
        self.update_clock()

    def update_word_count(self, event=None):
        text = self.input_text.get("1.0", 'end-1c')
        words = len(text.split())
        chars = len(text)
        self.input_word_count.config(text=f"📝 Words: {words} | Characters: {chars}")

    def update_clock(self):
        current_time = datetime.now().strftime("%H:%M:%S")
        self.time_label.config(text=f"🕐 {current_time}")
        self.root.after(1000, self.update_clock)

    def update_status(self, message, status_type="info"):
        icons = {"info": "✅", "warning": "⚠️", "error": "❌", "loading": "🔄"}
        icon = icons.get(status_type, "✅")
        self.status_label.config(text=f"{icon} {message}")
        self.root.update_idletasks()

    def get_text_and_count(self):
        text = self.input_text.get("1.0", 'end-1c').strip()
        if not text:
            messagebox.showwarning("Warning", "Please enter or upload text first!")
            return None, None

        count = self.sentence_count.get().strip()
        if not count.isdigit() or int(count) < 1:
            count = 5
            self.sentence_count.delete(0, tk.END)
            self.sentence_count.insert(0, "5")

        count = int(count)

        # Adjust if text is too short
        sentences = sent_tokenize(text)
        if count > len(sentences):
            count = max(1, len(sentences) // 2)
            self.update_status(f"Text is short, using {count} sentences", "warning")

        return text, count

    def lexrank_summary(self):
        """MOST IMPORTANT SENTENCES - LexRank Algorithm"""
        text, count = self.get_text_and_count()
        if not text:
            return

        self.update_status("Generating Most Important Sentences summary...", "loading")
        self.output_method.config(text="Method: 🎯 MOST IMPORTANT SENTENCES (LexRank)")

        try:
            parser = PlaintextParser.from_string(text, Tokenizer('english'))
            summarizer = LexRankSummarizer()
            summary_sentences = summarizer(parser.document, sentences_count=count)

            summary_text = self.format_summary(summary_sentences, "MOST IMPORTANT SENTENCES",
                                               "LexRank", "Finds central and most representative sentences")

            self.display_summary(summary_text, "LexRank", text)
            self.update_status("Most Important Sentences summary generated!", "info")

        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")
            self.update_status("Error generating summary", "error")

    def lsa_summary(self):
        """MAIN TOPIC SUMMARY - LSA Algorithm"""
        text, count = self.get_text_and_count()
        if not text:
            return

        self.update_status("Identifying main topics...", "loading")
        self.output_method.config(text="Method: 🔬 MAIN TOPIC SUMMARY (LSA)")

        try:
            parser = PlaintextParser.from_string(text, Tokenizer('english'))
            summarizer = LsaSummarizer()
            summary_sentences = summarizer(parser.document, sentences_count=count)

            summary_text = self.format_summary(summary_sentences, "MAIN TOPIC SUMMARY",
                                               "LSA", "Identifies and focuses on core themes")

            self.display_summary(summary_text, "LSA", text)
            self.update_status("Main Topic Summary generated!", "info")

        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")
            self.update_status("Error generating summary", "error")

    def luhn_summary(self):
        """QUICK KEY POINTS - Luhn Algorithm"""
        text, count = self.get_text_and_count()
        if not text:
            return

        self.update_status("Extracting key points...", "loading")
        self.output_method.config(text="Method: ⚡ QUICK KEY POINTS (Luhn)")

        try:
            parser = PlaintextParser.from_string(text, Tokenizer('english'))
            summarizer = LuhnSummarizer()
            summary_sentences = summarizer(parser.document, sentences_count=count)

            summary_text = self.format_summary(summary_sentences, "QUICK KEY POINTS",
                                               "Luhn", "Extracts important keywords and phrases")

            self.display_summary(summary_text, "Luhn", text)
            self.update_status("Quick Key Points generated!", "info")

        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")
            self.update_status("Error generating summary", "error")

    def kl_summary(self):
        """DETAILED IMPORTANT SUMMARY - KL Algorithm"""
        text, count = self.get_text_and_count()
        if not text:
            return

        self.update_status("Generating detailed summary...", "loading")
        self.output_method.config(text="Method: 📚 DETAILED IMPORTANT SUMMARY (KL)")

        try:
            parser = PlaintextParser.from_string(text, Tokenizer('english'))
            summarizer = KLSummarizer()
            summary_sentences = summarizer(parser.document, sentences_count=count)

            summary_text = self.format_summary(summary_sentences, "DETAILED IMPORTANT SUMMARY",
                                               "KL", "Preserves maximum information with full details")

            self.display_summary(summary_text, "KL", text)
            self.update_status("Detailed Important Summary generated!", "info")

        except Exception as e:
            messagebox.showerror("Error", f"Failed: {str(e)}")
            self.update_status("Error generating summary", "error")

    def format_summary(self, summary_sentences, title, method, description):
        """Format summary with professional styling"""
        style = self.summary_style.get()

        # Convert to text
        summary_text = " ".join(str(sentence) for sentence in summary_sentences)

        # Apply style modifications
        if style == "📝 Concise":
            words = summary_text.split()
            if len(words) > 150:
                summary_text = " ".join(words[:150]) + "..."
        elif style == "📖 Detailed":
            # Add more context
            original = self.input_text.get("1.0", 'end-1c')
            sentences = sent_tokenize(original)
            if len(summary_sentences) < 8 and len(sentences) > 15:
                summary_text = summary_text + " " + " ".join(sentences[-3:])

        # Create header
        header = f"""
╔{'═' * 70}╗
║ 🎯 {title} {' ' * (55 - len(title))}║
╠{'═' * 70}╣
║ 📊 Method: {method} {' ' * (55 - len(method))}║
║ 📝 {description} {' ' * (55 - len(description))}║
╠{'═' * 70}╣
║ ✨ Summary Statistics:                                    ║
║    • Sentences: {len(summary_sentences):<3}  |  Words: {len(summary_text.split()):<4}  ║
║    • Style: {style:<10}  |  Time: {datetime.now().strftime('%H:%M:%S')}        ║
╠{'═' * 70}╣
║ 📄 SUMMARY CONTENT:                                       ║
╠{'═' * 70}╣

"""
        footer = f"""
╚{'═' * 70}╝
💡 Tip: Try different methods to get various perspectives!
"""

        return header + summary_text.strip() + footer

    def display_summary(self, summary_text, method, original_text):
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(1.0, summary_text)

        # Save to history
        history_entry = {
            'id': len(self.history) + 1,
            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'method': method,
            'original_preview': original_text[:100] + "...",
            'summary_preview': summary_text[:100] + "...",
            'full_summary': summary_text
        }
        self.history.append(history_entry)
        self.save_history()

    def compare_all(self):
        """Compare all four methods"""
        text, count = self.get_text_and_count()
        if not text:
            return

        compare_window = tk.Toplevel(self.root)
        compare_window.title("📊 Compare All Summarization Methods")
        compare_window.geometry("1400x800")
        compare_window.configure(bg='#f5f5f5')

        tk.Label(compare_window, text="📊 METHOD COMPARISON", font=("Arial", 18, "bold"),
                 bg='#f5f5f5', fg='#2c3e50').pack(pady=10)

        notebook = ttk.Notebook(compare_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        methods_data = [
            ("🎯 Most Important Sentences", LexRankSummarizer, "#27ae60", "LexRank - Central sentences"),
            ("🔬 Main Topic Summary", LsaSummarizer, "#2980b9", "LSA - Core themes"),
            ("⚡ Quick Key Points", LuhnSummarizer, "#e67e22", "Luhn - Keywords"),
            ("📚 Detailed Important", KLSummarizer, "#8e44ad", "KL - Full information")
        ]

        parser = PlaintextParser.from_string(text, Tokenizer('english'))

        for title, SummarizerClass, color, subtitle in methods_data:
            frame = tk.Frame(notebook, bg='white')
            notebook.add(frame, text=title[:20])

            # Header
            header = tk.Frame(frame, bg=color, height=80)
            header.pack(fill=tk.X)
            header.pack_propagate(False)

            tk.Label(header, text=title, font=("Arial", 14, "bold"),
                     fg='white', bg=color).pack(pady=5)
            tk.Label(header, text=subtitle, font=("Arial", 10),
                     fg='white', bg=color).pack()

            # Summary
            text_widget = scrolledtext.ScrolledText(frame, wrap=tk.WORD, font=("Arial", 11))
            text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            try:
                summarizer = SummarizerClass()
                summary_sentences = summarizer(parser.document, sentences_count=count)

                result = f"{'=' * 60}\n"
                result += f"SUMMARY OUTPUT\n"
                result += f"{'=' * 60}\n\n"

                for i, sentence in enumerate(summary_sentences, 1):
                    result += f"{i}. {sentence}\n\n"

                result += f"\n{'=' * 60}\n"
                result += f"📊 Statistics:\n"
                result += f"• Sentences: {len(summary_sentences)}\n"
                result += f"• Words: {len(str(summary_sentences).split())}\n"
                result += f"• Characters: {len(str(summary_sentences))}\n"

                text_widget.insert(1.0, result)
                text_widget.config(state=tk.DISABLED)

            except Exception as e:
                text_widget.insert(1.0, f"Error: {str(e)}")

        self.update_status("Comparison window opened", "info")

    def upload_document(self):
        file_path = filedialog.askopenfilename(filetypes=[
            ("Text Files", "*.txt"), ("Word Documents", "*.docx"),
            ("PDF Files", "*.pdf"), ("All Files", "*.*")
        ])

        if file_path:
            try:
                content = ""
                if file_path.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as file:
                        content = file.read()
                elif file_path.endswith('.docx'):
                    doc = Document(file_path)
                    content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                elif file_path.endswith('.pdf'):
                    with open(file_path, 'rb') as file:
                        reader = PyPDF2.PdfReader(file)
                        content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

                self.input_text.delete(1.0, tk.END)
                self.input_text.insert(1.0, content)
                self.update_word_count()
                self.update_status(f"Loaded: {os.path.basename(file_path)}", "info")
                messagebox.showinfo("Success", "Document uploaded successfully!")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to read document: {str(e)}")

    def paste_text(self):
        try:
            text = self.root.clipboard_get()
            self.input_text.insert(tk.INSERT, text)
            self.update_word_count()
            self.update_status("Text pasted from clipboard", "info")
        except:
            messagebox.showwarning("Warning", "No text in clipboard")

    def clear_all(self):
        if messagebox.askyesno("Confirm", "Clear all text?"):
            self.input_text.delete(1.0, tk.END)
            self.output_text.delete(1.0, tk.END)
            self.update_word_count()
            self.output_method.config(text="Method: None")
            self.update_status("All text cleared", "info")

    def text_to_speech(self):
        text = self.output_text.get("1.0", 'end-1c').strip()
        if not text:
            messagebox.showwarning("Warning", "No summary to convert to audio")
            return

        # Remove formatting characters
        import re
        clean_text = re.sub(r'[╔╗║╠╣╚╝═]', '', text)
        clean_text = re.sub(r'[│┌┐└┘├┤┬┴┼]', '', clean_text)

        def speak():
            try:
                tts = gTTS(text=clean_text[:500], lang='en')
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp:
                    tts.save(temp.name)
                playsound.playsound(temp.name)
                os.unlink(temp.name)
            except Exception as e:
                self.root.after(0, lambda: messagebox.showerror("Error", f"Audio failed: {str(e)}"))

        threading.Thread(target=speak, daemon=True).start()
        self.update_status("Playing audio...", "info")

    def save_summary(self):
        summary = self.output_text.get("1.0", 'end-1c').strip()
        if not summary:
            messagebox.showwarning("Warning", "No summary to save")
            return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("Word Document", "*.docx"), ("All files", "*.*")]
        )

        if file_path:
            try:
                if file_path.endswith('.docx'):
                    doc = Document()
                    doc.add_heading('Text Summary', 0)
                    doc.add_paragraph(summary)
                    doc.save(file_path)
                else:
                    with open(file_path, 'w', encoding='utf-8') as file:
                        file.write(summary)

                self.update_status(f"Summary saved to {os.path.basename(file_path)}", "info")
                messagebox.showinfo("Success", "Summary saved successfully!")
            except Exception as e:
                messagebox.showerror("Error", f"Save failed: {str(e)}")

    def export_summary(self):
        self.save_summary()

    def add_to_favorites(self):
        summary = self.output_text.get("1.0", 'end-1c').strip()
        if not summary:
            messagebox.showwarning("Warning", "No summary to save to favorites")
            return

        name = tk.simpledialog.askstring("Favorite Name", "Enter a name for this summary:")
        if name:
            favorite = {
                'id': len(self.favorites) + 1,
                'name': name,
                'content': summary,
                'date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            self.favorites.append(favorite)
            self.save_favorites()
            self.update_status(f"Added '{name}' to favorites", "info")
            messagebox.showinfo("Success", "Added to favorites!")

    def open_favorites(self):
        if not self.favorites:
            messagebox.showinfo("Info", "No favorites yet. Save a summary to favorites first!")
            return

        fav_window = tk.Toplevel(self.root)
        fav_window.title("⭐ My Favorites")
        fav_window.geometry("600x500")
        fav_window.configure(bg='#f5f5f5')

        tk.Label(fav_window, text="⭐ SAVED FAVORITES", font=("Arial", 14, "bold"),
                 bg='#f5f5f5').pack(pady=10)

        listbox = tk.Listbox(fav_window, font=("Arial", 11), height=15)
        listbox.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        for fav in self.favorites:
            listbox.insert(tk.END, f"{fav['name']} - {fav['date']}")

        def load_favorite():
            selection = listbox.curselection()
            if selection:
                fav = self.favorites[selection[0]]
                self.output_text.delete(1.0, tk.END)
                self.output_text.insert(1.0, fav['content'])
                self.update_status(f"Loaded favorite: {fav['name']}", "info")
                fav_window.destroy()

        def delete_favorite():
            selection = listbox.curselection()
            if selection and messagebox.askyesno("Confirm", "Delete this favorite?"):
                del self.favorites[selection[0]]
                self.save_favorites()
                fav_window.destroy()
                self.open_favorites()

        btn_frame = tk.Frame(fav_window, bg='#f5f5f5')
        btn_frame.pack(pady=10)
        tk.Button(btn_frame, text="📖 Load", command=load_favorite, bg='#27ae60', fg='white').pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="🗑️ Delete", command=delete_favorite, bg='#e74c3c', fg='white').pack(side=tk.LEFT,
                                                                                                       padx=5)

    def open_history(self):
        if not self.history:
            messagebox.showinfo("Info", "No history yet. Generate some summaries first!")
            return

        history_window = tk.Toplevel(self.root)
        history_window.title("📜 History")
        history_window.geometry("900x600")
        history_window.configure(bg='#f5f5f5')

        tk.Label(history_window, text="📜 SUMMARY HISTORY", font=("Arial", 14, "bold"),
                 bg='#f5f5f5').pack(pady=10)

        # Create treeview
        columns = ('ID', 'Time', 'Method', 'Preview')
        tree = ttk.Treeview(history_window, columns=columns, show='headings', height=15)

        for col in columns:
            tree.heading(col, text=col)
            tree.column(col, width=150)

        tree.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        for item in self.history:
            tree.insert('', tk.END, values=(
                item['id'], item['timestamp'], item['method'], item['summary_preview'][:50]
            ))

        def load_history():
            selection = tree.selection()
            if selection:
                item = tree.item(selection[0])
                idx = int(item['values'][0]) - 1
                if 0 <= idx < len(self.history):
                    self.output_text.delete(1.0, tk.END)
                    self.output_text.insert(1.0, self.history[idx]['full_summary'])
                    self.update_status(f"Loaded history from {self.history[idx]['timestamp']}", "info")
                    history_window.destroy()

        tk.Button(history_window, text="📖 Load Selected", command=load_history,
                  bg='#27ae60', fg='white', font=("Arial", 11)).pack(pady=10)

    def open_translator(self):
        translator_window = tk.Toplevel(self.root)
        translator_window.title("🌐 Translator")
        translator_window.geometry("800x600")
        translator_window.configure(bg='#f5f5f5')

        tk.Label(translator_window, text="🌐 TEXT TRANSLATOR", font=("Arial", 14, "bold"),
                 bg='#f5f5f5').pack(pady=10)

        # Language selection
        lang_frame = tk.Frame(translator_window, bg='#f5f5f5')
        lang_frame.pack(pady=10)

        tk.Label(lang_frame, text="From:", font=("Arial", 11)).pack(side=tk.LEFT, padx=5)
        from_lang = ttk.Combobox(lang_frame, values=list(languages.keys()), width=15)
        from_lang.pack(side=tk.LEFT, padx=5)
        from_lang.set('English')

        tk.Label(lang_frame, text="To:", font=("Arial", 11)).pack(side=tk.LEFT, padx=5)
        to_lang = ttk.Combobox(lang_frame, values=list(languages.keys()), width=15)
        to_lang.pack(side=tk.LEFT, padx=5)
        to_lang.set('Hindi')

        # Text areas
        input_area = scrolledtext.ScrolledText(translator_window, wrap=tk.WORD,
                                               font=("Arial", 11), height=10)
        input_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        output_area = scrolledtext.ScrolledText(translator_window, wrap=tk.WORD,
                                                font=("Arial", 11), height=10)
        output_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        def translate():
            text = input_area.get("1.0", 'end-1c').strip()
            if not text:
                messagebox.showwarning("Warning", "Enter text to translate")
                return

            try:
                translator = GoogleTranslator(
                    source=languages[from_lang.get()],
                    target=languages[to_lang.get()]
                )
                translated = translator.translate(text)
                output_area.delete(1.0, tk.END)
                output_area.insert(1.0, translated)
                self.update_status(f"Translated from {from_lang.get()} to {to_lang.get()}", "info")
            except Exception as e:
                messagebox.showerror("Error", f"Translation failed: {str(e)}")

        tk.Button(translator_window, text="🌐 Translate", command=translate,
                  bg='#27ae60', fg='white', font=("Arial", 12, "bold")).pack(pady=10)

    def open_analyzer(self):
        analyzer_window = tk.Toplevel(self.root)
        analyzer_window.title("📊 Text Analyzer")
        analyzer_window.geometry("800x600")
        analyzer_window.configure(bg='#f5f5f5')

        tk.Label(analyzer_window, text="📊 TEXT ANALYZER", font=("Arial", 14, "bold"),
                 bg='#f5f5f5').pack(pady=10)

        text_area = scrolledtext.ScrolledText(analyzer_window, wrap=tk.WORD,
                                              font=("Arial", 11), height=10)
        text_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        result_area = scrolledtext.ScrolledText(analyzer_window, wrap=tk.WORD,
                                                font=("Arial", 11), height=10)
        result_area.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        def analyze():
            text = text_area.get("1.0", 'end-1c').strip()
            if not text:
                messagebox.showwarning("Warning", "Enter text to analyze")
                return

            sentences = sent_tokenize(text)
            words = text.split()

            analysis = f"""
📊 TEXT ANALYSIS REPORT
{'=' * 50}

📝 Basic Statistics:
• Total Characters: {len(text)}
• Total Words: {len(words)}
• Total Sentences: {len(sentences)}
• Average Word Length: {sum(len(w) for w in words) / len(words):.2f}
• Average Sentence Length: {len(words) / len(sentences):.2f} words

📈 Readability:
• Flesch Reading Ease: {206.835 - 1.015 * (len(words) / len(sentences)) - 84.6 * (len(words) / len(sentences)):.2f}

🔑 Top Keywords:
"""
            # Simple keyword extraction
            from collections import Counter
            stopwords = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is',
                         'are', 'was', 'were'}
            keywords = [w.lower() for w in words if w.lower() not in stopwords and len(w) > 3]
            for word, count in Counter(keywords).most_common(10):
                analysis += f"• {word}: {count} times\n"

            result_area.delete(1.0, tk.END)
            result_area.insert(1.0, analysis)

        tk.Button(analyzer_window, text="🔍 Analyze Text", command=analyze,
                  bg='#27ae60', fg='white', font=("Arial", 12, "bold")).pack(pady=10)

    def open_batch_processor(self):
        messagebox.showinfo("Info", "Batch Processing feature coming soon!")

    def change_theme(self):
        messagebox.showinfo("Info", "Theme customization coming soon!")

    def audio_settings(self):
        messagebox.showinfo("Info", "Audio settings coming soon!")

    def clear_cache(self):
        temp_dir = tempfile.gettempdir()
        cleaned = 0
        for file in os.listdir(temp_dir):
            if file.endswith('.mp3') and file.startswith('tmp'):
                try:
                    os.remove(os.path.join(temp_dir, file))
                    cleaned += 1
                except:
                    pass
        messagebox.showinfo("Cache Cleared", f"Cleaned {cleaned} temporary files")
        self.update_status(f"Cache cleared: {cleaned} files", "info")

    def new_session(self):
        if messagebox.askyesno("New Session", "Start new session? Current data will be cleared."):
            self.input_text.delete(1.0, tk.END)
            self.output_text.delete(1.0, tk.END)
            self.update_word_count()
            self.output_method.config(text="Method: None")
            self.update_status("New session started", "info")

    def show_documentation(self):
        doc_window = tk.Toplevel(self.root)
        doc_window.title("Documentation")
        doc_window.geometry("700x600")
        doc_window.configure(bg='#f5f5f5')

        doc_text = scrolledtext.ScrolledText(doc_window, wrap=tk.WORD, font=("Arial", 11))
        doc_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        documentation = """
📚 USER MANUAL - Text Summarization Suite v3.0

============================================

🎯 SUMMARIZATION METHODS:

1. MOST IMPORTANT SENTENCES (LexRank)
   • Best for: News articles, blogs, general text
   • How it works: Finds central sentences using graph algorithm
   • Output: Balanced representation of main ideas

2. MAIN TOPIC SUMMARY (LSA)
   • Best for: Scientific papers, technical documents
   • How it works: Identifies hidden topics using SVD
   • Output: Topic-focused summary

3. QUICK KEY POINTS (Luhn)
   • Best for: Product descriptions, keyword-rich text
   • How it works: Scores sentences by keyword frequency
   • Output: Keyword-focused bullet points

4. DETAILED IMPORTANT SUMMARY (KL)
   • Best for: Legal documents, critical information
   • How it works: Minimizes information loss
   • Output: Comprehensive detailed summary

============================================

✨ FEATURES:

• Upload files (TXT, DOCX, PDF)
• Copy/Paste text directly
• Compare all methods side-by-side
• Text-to-speech audio output
• Save summaries to file
• Add favorites for quick access
• View history of all summaries
• Multi-language translation
• Text analysis tools
• Professional status bar

============================================

💡 TIPS:

• Use longer texts (1000+ words) for better results
• Try different methods to get various perspectives
• Use Compare All feature to see differences
• Save important summaries to favorites
• Clear cache periodically for better performance

============================================

Shortcuts:
• Ctrl+V: Paste text
• Ctrl+C: Copy text
• Ctrl+A: Select all

For support: contact@example.com
"""
        doc_text.insert(1.0, documentation)
        doc_text.config(state=tk.DISABLED)

    def show_about(self):
        about_window = tk.Toplevel(self.root)
        about_window.title("About")
        about_window.geometry("500x400")
        about_window.configure(bg='#f5f5f5')

        about_text = """
🎯 Professional Text Summarization Suite
Version 3.0

A comprehensive text processing tool with
advanced summarization algorithms.

✨ Features:
• 4 Advanced Summarization Methods
• Multi-language Translation
• Text Analysis Tools
• History & Favorites
• Audio Output
• Professional UI

📊 Algorithms:
• LexRank (Graph-based)
• LSA (Topic-based)
• Luhn (Keyword-based)
• KL (Information-based)

Developed with:
Python • Tkinter • NLTK • Sumy

© 2024 All Rights Reserved
"""
        tk.Label(about_window, text=about_text, font=("Arial", 11),
                 justify=tk.LEFT, bg='#f5f5f5').pack(padx=20, pady=20)
        tk.Button(about_window, text="Close", command=about_window.destroy,
                  bg='#27ae60', fg='white', font=("Arial", 10, "bold")).pack(pady=10)

    def play_welcome(self):
        def welcome():
            try:
                tts = gTTS(text="Welcome to Professional Text Summarization Suite", lang='en')
                with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp:
                    tts.save(temp.name)
                playsound.playsound(temp.name)
                os.unlink(temp.name)
            except:
                pass

        threading.Thread(target=welcome, daemon=True).start()

    def save_history(self):
        try:
            with open('history.json', 'w', encoding='utf-8') as f:
                json.dump(self.history, f, ensure_ascii=False, indent=2)
        except:
            pass

    def load_history(self):
        try:
            with open('history.json', 'r', encoding='utf-8') as f:
                self.history = json.load(f)
        except:
            self.history = []

    def save_favorites(self):
        try:
            with open('favorites.json', 'w', encoding='utf-8') as f:
                json.dump(self.favorites, f, ensure_ascii=False, indent=2)
        except:
            pass

    def load_favorites(self):
        try:
            with open('favorites.json', 'r', encoding='utf-8') as f:
                self.favorites = json.load(f)
        except:
            self.favorites = []

    def load_settings(self):
        try:
            with open('settings.json', 'r') as f:
                settings = json.load(f)
                self.sentence_count.insert(0, settings.get('sentence_count', '5'))
                self.summary_style.set(settings.get('summary_style', '⚖️ Balanced'))
        except:
            pass


# Run the application
if __name__ == "__main__":
    root = tk.Tk()
    app = ProfessionalSummarizer(root)
    root.mainloop()
